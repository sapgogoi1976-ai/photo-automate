from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from tkinter import Tk, filedialog, messagebox

def insert_photos_repeated(filename_list, repeats_per_line=6):
    # Create a new Word document
    doc = Document()
    
    # Set margins to minimum
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.5)  # Reduce top margin
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

    for filename in filename_list:
        # Insert the image repeated across one line
        table = doc.add_table(rows=1, cols=repeats_per_line)
        table.autofit = False
        table.allow_autofit = False
        
        # Set table properties to reduce spacing
        tbl = table._element
        tblPr = tbl.tblPr
        
        # Remove spacing before and after table
        tblSpacing = OxmlElement('w:tblSpacing')
        tblSpacing.set(qn('w:w'), '0')
        tblSpacing.set(qn('w:type'), 'nil')
        tblPr.append(tblSpacing)
        
        row = table.rows[0]
        row.height = Cm(3.6)  # Set row height to 3.6 cm
        
        # Set row properties
        trPr = row._element.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '2038')  # Height in twips (3.6 cm)
        trHeight.set(qn('w:type'), 'atLeast')
        trPr.append(trHeight)
        
        for i in range(repeats_per_line):
            cell = row.cells[i]
            
            # Set cell margins to zero
            tcPr = cell._element.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            
            for margin_name in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{margin_name}')
                margin.set(qn('w:w'), '0')
                margin.set(qn('w:type'), 'nil')
                tcMar.append(margin)
            
            tcPr.append(tcMar)
            
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Remove all paragraph spacing
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing = 1.0
            
            run = paragraph.add_run()
            # Set photo size to 3 cm width and 3.6 cm height
            run.add_picture(filename, width=Cm(3), height=Cm(3.6))
        
        # Add minimal spacing paragraph with very small gap
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(0)
        spacing_para.paragraph_format.space_after = Pt(1)  # Reduced to 1 point
        spacing_para.paragraph_format.line_spacing = 1.0
    
    # Save the document
    doc.save('Repeated_Photos.docx')

def select_photos():
    # Use Tkinter to open a file dialog to select photos
    root = Tk()
    root.withdraw()  # Hide the root window
    file_paths = filedialog.askopenfilenames(
        title="Select Photos",
        filetypes=[("Image Files", "*.png;*.jpg;*.jpeg;*.bmp;*.gif")]
    )
    return root.tk.splitlist(file_paths)

if __name__ == "__main__":
    photos = select_photos()
    if photos:
        insert_photos_repeated(photos)
        
        # Create a hidden root window to show message box
        root = Tk()
        root.withdraw()
        messagebox.showinfo("Success", "Word document created as 'Repeated_Photos.docx'")
        root.destroy()
    else:
        print("No photos were selected.")
